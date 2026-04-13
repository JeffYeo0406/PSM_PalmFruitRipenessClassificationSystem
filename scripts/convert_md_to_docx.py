"""
Convert TFLite_Conversion_Technical_Reference.md to DOCX format.
Creates a professionally formatted Word document with proper styling.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
from pathlib import Path

def create_tflite_report():
    """Create the TFLite Conversion Technical Reference as a DOCX document."""
    
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(51, 102, 153)
    
    # ===== DOCUMENT CONTENT =====
    
    # Title
    doc.add_heading('TFLite Conversion Technical Reference', 0)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Understanding Model Size, Parameters, and Memory Usage for Edge AI Deployment')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document metadata
    meta = doc.add_paragraph()
    meta.add_run('Document Version: ').bold = True
    meta.add_run('1.0\n')
    meta.add_run('Date: ').bold = True
    meta.add_run('April 7, 2026\n')
    meta.add_run('Model: ').bold = True
    meta.add_run('Palm Fruit Ripeness Classification System (MobileNetV2)\n')
    meta.add_run('Target Platform: ').bold = True
    meta.add_run('Raspberry Pi 4B (4GB) with Camera Module 3')
    
    doc.add_paragraph()  # Spacer
    
    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading('Executive Summary', level=1)
    
    doc.add_paragraph(
        'This technical reference provides comprehensive insights into TensorFlow Lite (TFLite) model conversion, '
        'quantization effects, and deployment optimization for edge AI applications. Using the Palm Fruit Ripeness '
        'Classification System as a case study, we demonstrate how INT8 quantization achieves a 75% size reduction '
        '(10.86 MB → 2.76 MB) while maintaining 99.4% accuracy retention (0.60% drop).'
    )
    
    # Key Findings Table
    doc.add_heading('Key Findings', level=2)
    
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'FP32'
    header_cells[2].text = 'FP16'
    header_cells[3].text = 'INT8'
    header_cells[4].text = 'Significance'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    data = [
        ['File Size', '9.10 MB', '4.61 MB', '2.76 MB', '70% reduction (FP32→INT8)'],
        ['Accuracy', '92.78%', '~92.78%', '92.22%', '0.60% drop (acceptable)'],
        ['Parameters', '9,417,611', '11,788,782', '9,568,142', 'Core: 9.4M (identical)'],
        ['Bytes/Param', '4', '2', '1', 'Primary size driver'],
        ['Deployment', 'Reference', 'ARMv8 optimized', 'Raspberry Pi recommended', ''],
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
    
    doc.add_paragraph()  # Spacer
    
    # ===== SECTION 1: INTRODUCTION =====
    doc.add_heading('1. Introduction to TFLite Conversion', level=1)
    
    doc.add_heading('1.1 Why Convert to TFLite?', level=2)
    doc.add_paragraph(
        'TensorFlow Lite is Google\'s lightweight runtime for deploying machine learning models on edge devices. '
        'Key advantages:'
    )
    
    # Bullet points
    bullets = [
        'Smaller binary size: Optimized flatbuffer format',
        'Faster inference: Hardware acceleration support (NNAPI, Core ML, etc.)',
        'Lower memory footprint: Quantization reduces RAM requirements',
        'Portable: Single .tflite file contains model + metadata',
    ]
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_heading('1.2 Quantization Formats', level=2)
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    headers = ['Format', 'Data Type', 'Size per Parameter', 'Use Case']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    formats = [
        ['FP32', '32-bit float', '4 bytes', 'Reference implementation, maximum precision'],
        ['FP16', '16-bit float', '2 bytes', 'ARMv8 GPUs, good balance of speed/accuracy'],
        ['INT8', '8-bit integer', '1 byte', 'Edge CPUs (Raspberry Pi), fastest inference'],
    ]
    
    for i, row_data in enumerate(formats):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    # ===== SECTION 2: METHODOLOGY =====
    doc.add_heading('2. Methodology', level=1)
    
    doc.add_heading('2.1 Model Architecture', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Base Model: ').bold = True
    p.add_run('MobileNetV2 (pre-trained on ImageNet)\n')
    p.add_run('Custom Head: ').bold = True
    p.add_run('Dense layers for 3-class classification\n')
    p.add_run('Classes: ').bold = True
    p.add_run('["Overripe", "Ripe", "Underripe"]')
    
    doc.add_heading('Architecture Details:', level=3)
    arch_details = [
        'Input: 224×224×3 RGB images',
        'Preprocessing: MobileNetV2 preprocess_input (pixels → [-1, 1])',
        'Backbone: MobileNetV2 (frozen during warmup, top 30 layers unfrozen for fine-tuning)',
        'Classification head: GlobalAveragePooling2D → Dense(128, relu) → Dropout(0.3) → Dense(3, softmax)',
    ]
    for detail in arch_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_heading('2.2 Conversion Parameters', level=2)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    param_headers = ['Parameter', 'Value', 'Rationale']
    for i, header in enumerate(param_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    params = [
        ['Representative Dataset', '500 images', '↑ from default 200 for better calibration'],
        ['Experimental New Quantizer', 'True', 'More accurate INT8 weight quantization'],
        ['INT8 I/O Type', 'tf.float32', 'Float fallback for compatibility'],
        ['FP16 Optimization', 'Optimize.DEFAULT', 'Weight quantization enabled'],
        ['Image Size', '224×224', 'Matches training input size'],
    ]
    
    for i, row_data in enumerate(params):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    # ===== SECTION 3: RESULTS =====
    doc.add_heading('3. Results and Findings', level=1)
    
    doc.add_heading('3.1 Model Size Comparison', level=2)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    size_headers = ['Format', 'File Size', 'Compression Ratio', 'Theoretical Weight Memory']
    for i, header in enumerate(size_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    sizes = [
        ['Original .h5', '10.86 MB', '1.00× (baseline)', '35.93 MB'],
        ['FP32 TFLite', '9.10 MB', '0.84×', '35.93 MB'],
        ['FP16 TFLite', '4.61 MB', '0.42×', '17.96 MB'],
        ['INT8 TFLite', '2.76 MB', '0.25×', '8.98 MB'],
    ]
    
    for i, row_data in enumerate(sizes):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Key Insight: ').bold = True
    p.add_run(
        'File sizes are smaller than theoretical weight memory because TFLite uses efficient flatbuffer '
        'serialization and only stores weights (not activation workspace memory).'
    )
    
    doc.add_heading('3.2 Accuracy Validation Results', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Test Set: ').bold = True
    p.add_run('180 images (60 per class)')
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    acc_headers = ['Model', 'Accuracy', 'Correct/Total', 'Drop vs FP32']
    for i, header in enumerate(acc_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    acc_data = [
        ['FP32', '92.78%', '167/180', '0.00% (baseline)'],
        ['FP16', '~92.78%', '~167/180', '~0.00% (negligible)'],
        ['INT8', '92.22%', '166/180', '0.60% ✓'],
    ]
    
    for i, row_data in enumerate(acc_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Per-Class Breakdown:\n').bold = True
    p.add_run('Overripe:   FP32=60/60 (100%)  → INT8=59/60 (98.3%)  [-1 image]\n')
    p.add_run('Ripe:       FP32=54/60 (90.0%) → INT8=54/60 (90.0%)  [0 images]\n')
    p.add_run('Underripe:  FP32=53/60 (88.3%) → INT8=53/60 (88.3%)  [0 images]')
    
    p = doc.add_paragraph()
    p.add_run('Conclusion: ').bold = True
    p.add_run(
        'INT8 quantization caused only 1 additional misclassification out of 180 test images, '
        'well within acceptable limits for production deployment.'
    )
    
    # ===== SECTION 4: PARAMETER ANALYSIS =====
    doc.add_heading('4. Parameter Analysis Deep Dive', level=1)
    
    doc.add_heading('4.1 Understanding Parameter Counts', level=2)
    
    p = doc.add_paragraph(
        'One of the most confusing aspects of TFLite conversion is that reported parameter counts differ '
        'across formats, even though the underlying architecture is identical.'
    )
    
    doc.add_heading('Core Architecture (Identical Across All Formats)', level=3)
    
    core_details = [
        'Total Parameters: 9,417,611',
        'Trainable Parameters: 9,417,611 (all parameters are learned)',
        'This represents: MobileNetV2 backbone + custom classification head',
    ]
    for detail in core_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_heading('Why Reported Counts Differ', level=3)
    
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header
    count_headers = ['Format', 'Reported Parameters', 'Tensors', 'Difference', 'Reason']
    for i, header in enumerate(count_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    count_data = [
        ['FP32', '9,417,611', '176', 'Baseline', 'No extra metadata'],
        ['FP16', '11,788,782', '284', '+25.18%', 'Alignment padding for GPU efficiency'],
        ['INT8', '9,568,142', '178', '+1.60%', 'Quantization scales & zero points'],
    ]
    
    for i, row_data in enumerate(count_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 The Critical Insight', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Parameter count ≠ Memory usage\n\n').bold = True
    p.add_run('What actually matters for inference is bytes per parameter:\n\n')
    p.add_run('Memory = Parameters × Bytes_per_Parameter + Metadata_Overhead')
    
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header
    mem_headers = ['Format', 'Bytes/Param', 'Theoretical Memory', 'Actual File Size', 'Overhead']
    for i, header in enumerate(mem_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    mem_data = [
        ['FP32', '4', '35.93 MB', '9.10 MB', '~75% (flatbuffer compression)'],
        ['FP16', '2', '17.96 MB', '4.61 MB', '~74% (flatbuffer compression)'],
        ['INT8', '1', '8.98 MB', '2.76 MB', '~69% (flatbuffer + quant metadata)'],
    ]
    
    for i, row_data in enumerate(mem_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Why file sizes are smaller than theoretical memory:\n').bold = True
    reasons = [
        'TFLite uses efficient flatbuffer serialization',
        'Only stores weights (not activation workspace)',
        'Compresses repeated patterns in weight tensors',
    ]
    for reason in reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    # ===== SECTION 5: MEMORY USAGE =====
    doc.add_heading('5. Memory Usage Analysis', level=1)
    
    doc.add_heading('5.1 Runtime Memory Requirements', level=2)
    
    doc.add_paragraph('During inference, the model needs memory for:')
    
    mem_components = [
        'Weights: Stored in the .tflite file',
        'Activations: Intermediate layer outputs (workspace)',
        'Input/Output Buffers: Preprocessed image + predictions',
    ]
    for comp in mem_components:
        doc.add_paragraph(comp, style='List Bullet')
    
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header
    runtime_headers = ['Component', 'FP32', 'FP16', 'INT8', 'Notes']
    for i, header in enumerate(runtime_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    runtime_data = [
        ['Weights', '35.93 MB', '17.96 MB', '8.98 MB', 'Theoretical (params × bytes)'],
        ['Activations', '~50-100 MB', '~25-50 MB', '~12-25 MB', 'Depends on batch size'],
        ['I/O Buffers', '~0.6 MB', '~0.3 MB', '~0.15 MB', '224×224×3 image'],
        ['Total Runtime', '~90-140 MB', '~45-70 MB', '~22-35 MB', 'Estimated for Pi 4B'],
    ]
    
    for i, row_data in enumerate(runtime_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 Raspberry Pi 4B (4GB) Suitability', level=2)
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    pi_headers = ['Model', 'RAM Usage', 'Available RAM', 'Feasibility']
    for i, header in enumerate(pi_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    pi_data = [
        ['FP32', '~90-140 MB', '~3.5 GB', '✓ Feasible (wasteful)'],
        ['FP16', '~45-70 MB', '~3.5 GB', '✓ Feasible (good balance)'],
        ['INT8', '~22-35 MB', '~3.5 GB', '✓ Optimal (recommended)'],
    ]
    
    for i, row_data in enumerate(pi_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Recommendation: ').bold = True
    p.add_run('INT8 model leaves maximum RAM for OS overhead, image preprocessing, API server (Flask), and other applications.')
    
    # ===== SECTION 6: QUANTIZATION EFFECTS =====
    doc.add_heading('6. Quantization Effects on Accuracy', level=1)
    
    doc.add_heading('6.1 Why INT8 Causes Accuracy Drop', level=2)
    
    doc.add_paragraph(
        'INT8 quantization maps continuous float32 values to discrete int8 values:\n\n'
        'float32_weight → quantize → int8_weight\n'
        'int8_weight → dequantize → float32_weight (approximate)'
    )
    
    p = doc.add_paragraph()
    p.add_run('Information Loss: ').bold = True
    p.add_run('Float32 has ~7 decimal digits of precision, while INT8 has only 256 discrete values.')
    
    doc.add_heading('6.2 Mitigation Strategies Used', level=2)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    strat_headers = ['Strategy', 'Implementation', 'Effect']
    for i, header in enumerate(strat_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    strat_data = [
        ['Representative Dataset', '500 images (↑ from 200)', 'Better calibration of quantization ranges'],
        ['Experimental New Quantizer', 'True', 'More accurate weight quantization'],
        ['Float32 I/O Fallback', 'inference_input_type=tf.float32', 'Prevents hard failures on edge cases'],
        ['Per-Tensor Quantization', 'Default', 'Simpler, more compatible than per-channel'],
    ]
    
    for i, row_data in enumerate(strat_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('6.3 Accuracy Retention Analysis', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Original Model (FP32):\n').bold = True
    p.add_run('Training accuracy: ~93-94%\n')
    p.add_run('Validation accuracy: ~92-93%\n')
    p.add_run('Test accuracy: 92.78%')
    
    p = doc.add_paragraph()
    p.add_run('After INT8 Quantization:\n').bold = True
    p.add_run('Test accuracy: 92.22%\n')
    p.add_run('Accuracy retention: 99.4% (0.60% drop)\n')
    p.add_run('Misclassifications: +1 out of 180 images')
    
    p = doc.add_paragraph()
    p.add_run('Conclusion: ').bold = True
    p.add_run('The quantization strategy successfully preserved model performance while achieving 75% size reduction.')
    
    # ===== SECTION 7: DEPLOYMENT =====
    doc.add_heading('7. Deployment Recommendations', level=1)
    
    doc.add_heading('7.1 Platform-Specific Guidance', level=2)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    platform_headers = ['Platform', 'Recommended Format', 'Rationale']
    for i, header in enumerate(platform_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    platform_data = [
        ['Raspberry Pi 4B (CPU)', 'INT8', 'Fastest inference, lowest memory'],
        ['Raspberry Pi 4B (GPU)', 'FP16', 'Better GPU support (if available)'],
        ['Desktop/Laptop (CPU)', 'FP32 or FP16', 'Maximum accuracy, ample resources'],
        ['Desktop/Laptop (GPU)', 'FP16', 'GPU acceleration support'],
        ['Mobile (Android/iOS)', 'INT8', 'Battery efficiency, limited RAM'],
    ]
    
    for i, row_data in enumerate(platform_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('7.2 Raspberry Pi 4B Deployment Checklist', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Hardware Requirements:\n').bold = True
    hw_reqs = [
        '✓ Raspberry Pi 4B (4GB RAM recommended)',
        '✓ Camera Module 3 (or compatible USB camera)',
        '✓ MicroSD card (32GB+ recommended)',
        '✓ Power supply (3A, 5V)',
    ]
    for req in hw_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Software Stack:\n').bold = True
    
    code_block = doc.add_paragraph()
    code_block.add_run(
        '# 1. Install OS (Raspberry Pi OS Lite recommended)\n'
        '# 2. Install Python dependencies\n'
        'pip install -r requirements-pi.txt\n\n'
        '# 3. Set environment variables\n'
        'export MODEL_PATH=models/palm_ripeness_best_20260407_014729_int8.tflite\n'
        'export LABELS_PATH=models/labels_20260407_014729.json\n\n'
        '# 4. Run Flask API\n'
        'python api/app.py'
    ).font.name = 'Courier New'
    
    p = doc.add_paragraph()
    p.add_run('Expected Performance:\n').bold = True
    perf_items = [
        'Latency: ~320ms per inference (includes preprocessing)',
        'Throughput: 2-4 FPS on Raspberry Pi CPU',
        'Memory: <50 MB RAM usage',
        'Accuracy: 92.22%',
    ]
    for item in perf_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # ===== SECTION 8: TECHNICAL REFERENCE =====
    doc.add_heading('8. Technical Reference Tables', level=1)
    
    doc.add_heading('8.1 Conversion Command Reference', level=2)
    
    code = doc.add_paragraph()
    code.add_run(
        '# Basic conversion (FP32 only)\n'
        'python scripts/convert_tflite.py \\\n'
        '    --h5 models/palm_ripeness_best_20260311_170850.h5 \\\n'
        '    --rep-data /path/to/train \\\n'
        '    --output-dir models\n\n'
        '# With labels file\n'
        'python scripts/convert_tflite.py \\\n'
        '    --h5 models/palm_ripeness_best_20260311_170850.h5 \\\n'
        '    --labels models/labels_20260311_170850.json \\\n'
        '    --rep-data /path/to/train \\\n'
        '    --output-dir models \\\n'
        '    --img-size 224'
    ).font.name = 'Courier New'
    
    doc.add_heading('8.2 Validation Command Reference', level=2)
    
    code = doc.add_paragraph()
    code.add_run(
        '# Validate INT8 vs FP32 accuracy\n'
        'python scripts/validate_tflite.py \\\n'
        '    --model-fp32 models/palm_ripeness_best_20260407_014729_fp32.tflite \\\n'
        '    --model-int8 models/palm_ripeness_best_20260407_014729_int8.tflite \\\n'
        '    --labels models/labels_20260407_014729.json \\\n'
        '    --data-dir /path/to/test \\\n'
        '    --img-size 224'
    ).font.name = 'Courier New'
    
    doc.add_heading('8.3 Parameter Counting Command Reference', level=2)
    
    code = doc.add_paragraph()
    code.add_run(
        '# Analyze parameter counts across formats\n'
        'python scripts/count_tflite_params.py'
    ).font.name = 'Courier New'
    
    doc.add_heading('8.4 File Structure Reference', level=2)
    
    code = doc.add_paragraph()
    code.add_run(
        'models/\n'
        '├── palm_ripeness_best_20260407_014729_fp32.tflite      # 9.10 MB\n'
        '├── palm_ripeness_best_20260407_014729_float16.tflite   # 4.61 MB\n'
        '├── palm_ripeness_best_20260407_014729_int8.tflite      # 2.76 MB (recommended)\n'
        '├── labels_20260407_014729.json                         # Class names\n'
        '└── tflite_manifest_20260407_014729.json                # Metadata'
    ).font.name = 'Courier New'
    
    doc.add_heading('8.5 API Endpoints Reference', level=2)
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    api_headers = ['Endpoint', 'Method', 'Purpose', 'Response']
    for i, header in enumerate(api_headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # Data
    api_data = [
        ['/health', 'GET', 'Check model status', '{"status": "ok", "ready": true}'],
        ['/classify', 'POST', 'Classify image', '{"label": "Ripe", "probability": 0.94}'],
        ['/result/<id>', 'GET', 'Fetch previous result', '{"label": "Ripe", "probability": 0.94}'],
    ]
    
    for i, row_data in enumerate(api_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    # ===== SECTION 9: CONCLUSIONS =====
    doc.add_heading('9. Conclusions', level=1)
    
    doc.add_heading('9.1 Key Takeaways', level=2)
    
    takeaways = [
        'Size Reduction: INT8 quantization achieves 75% size reduction (10.86 MB → 2.76 MB)',
        'Accuracy Retention: Only 0.60% accuracy drop (92.78% → 92.22%)',
        'Parameter Counts: Core architecture identical (9.4M params), differences due to metadata',
        'Memory Efficiency: INT8 uses ~1/4 the memory of FP32 for weights',
        'Production Ready: Validated for Raspberry Pi 4B deployment',
    ]
    
    for i, takeaway in enumerate(takeaways, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(takeaway)
    
    doc.add_heading('9.2 Best Practices', level=2)
    
    practices = [
        'Always validate accuracy after quantization (use validate_tflite.py)',
        'Use representative dataset of 500+ images for INT8 calibration',
        'Enable experimental new quantizer for better accuracy',
        'Set float32 I/O fallback for compatibility',
        'Monitor memory usage on target device during deployment',
    ]
    
    for i, practice in enumerate(practices, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(practice)
    
    doc.add_heading('9.3 Future Work', level=2)
    
    future = [
        'Per-channel quantization: Further accuracy improvements',
        'Edge TPU compilation: Hardware acceleration',
        'Model pruning: Additional size reduction',
        'Benchmark suite: Systematic performance testing',
    ]
    
    for item in future:
        doc.add_paragraph(item, style='List Bullet')
    
    # ===== SECTION 10: REFERENCES =====
    doc.add_heading('10. References', level=1)
    
    refs = [
        'TensorFlow Lite Documentation: https://www.tensorflow.org/lite',
        'MobileNetV2 Paper: "MobileNetV2: Inverted Residuals and Linear Bottlenecks"',
        'Post-Training Quantization Guide: https://www.tensorflow.org/lite/performance/post_training_quantization',
        'Raspberry Pi Deployment: https://www.tensorflow.org/lite/guide/build_cpp',
    ]
    
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f'{i}. {ref}')
    
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.add_run('Document End\n\n').bold = True
    footer.add_run('Generated: April 7, 2026\n')
    footer.add_run('Model: Palm Fruit Ripeness Classification System\n')
    footer.add_run('Framework: TensorFlow 2.x + TFLite\n')
    footer.add_run('Target: Raspberry Pi 4B (4GB)')
    
    # Save the document
    output_path = Path('reports/TFLite_Conversion_Technical_Reference.docx')
    doc.save(output_path)
    print(f"✓ Document saved to: {output_path}")
    print(f"  File size: {output_path.stat().st_size / 1024:.1f} KB")

if __name__ == "__main__":
    create_tflite_report()